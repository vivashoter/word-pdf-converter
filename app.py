from flask import Flask, request, send_file, send_from_directory, jsonify
from werkzeug.utils import secure_filename

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import fitz
import os
import re
import subprocess
import tempfile


# ============================================================
# CONFIG
# ============================================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__)

app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024


# ============================================================
# STATIC ROUTES
# ============================================================

@app.route("/")
def home():
    return send_from_directory(BASE_DIR, "index.html")


@app.route("/privacy.html")
def privacy():
    return send_from_directory(BASE_DIR, "privacy.html")


@app.route("/terms.html")
def terms():
    return send_from_directory(BASE_DIR, "terms.html")


@app.route("/about.html")
def about():
    return send_from_directory(BASE_DIR, "about.html")


@app.route("/contact.html")
def contact():
    return send_from_directory(BASE_DIR, "contact.html")


@app.route("/style.css")
def styles():
    return send_from_directory(BASE_DIR, "style.css")


@app.route("/script.js")
def scripts():
    return send_from_directory(BASE_DIR, "script.js")


@app.route("/convertdocgoose-logo.png")
def logo():
    return send_from_directory(BASE_DIR, "convertdocgoose-logo.png")


@app.route("/favicon.png")
def favicon():
    return send_from_directory(BASE_DIR, "favicon.png")


@app.errorhandler(413)
def file_too_large(error):
    return jsonify({
        "error": "File is too large. Maximum size is 25 MB."
    }), 413


# ============================================================
# WORD -> PDF
# ============================================================

@app.route("/convert/word-to-pdf", methods=["POST"])
def word_to_pdf():

    if "file" not in request.files:
        return jsonify({
            "error": "No file uploaded."
        }), 400

    uploaded_file = request.files["file"]

    if uploaded_file.filename == "":
        return jsonify({
            "error": "No file selected."
        }), 400

    filename = secure_filename(
        uploaded_file.filename
    )

    if not filename.lower().endswith(
        (".doc", ".docx")
    ):
        return jsonify({
            "error": "Please upload a DOC or DOCX file."
        }), 400

    with tempfile.TemporaryDirectory() as temp_dir:

        input_path = os.path.join(
            temp_dir,
            filename
        )

        uploaded_file.save(
            input_path
        )

        try:

            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    temp_dir,
                    input_path,
                ],
                check=True,
                timeout=120,
                capture_output=True,
                text=True,
            )

            if result.stdout:
                print(
                    "LibreOffice:",
                    result.stdout,
                    flush=True
                )

            if result.stderr:
                print(
                    "LibreOffice warnings:",
                    result.stderr,
                    flush=True
                )

        except subprocess.TimeoutExpired:

            return jsonify({
                "error": "The conversion took too long."
            }), 504

        except subprocess.CalledProcessError as error:

            print(
                "LibreOffice error:",
                error.stderr,
                flush=True
            )

            return jsonify({
                "error":
                "The Word document could not be converted."
            }), 500

        except Exception as error:

            print(
                "Word to PDF error:",
                repr(error),
                flush=True
            )

            return jsonify({
                "error":
                "The Word document could not be converted."
            }), 500

        output_name = (
            os.path.splitext(filename)[0]
            + ".pdf"
        )

        output_path = os.path.join(
            temp_dir,
            output_name
        )

        if not os.path.exists(
            output_path
        ):
            return jsonify({
                "error":
                "The PDF could not be created."
            }), 500

        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_name,
            mimetype="application/pdf"
        )


# ============================================================
# EXACTDOC
# ============================================================

def run_exactdoc(
    input_path,
    output_path
):

    print(
        "Running ExactDoc on:",
        input_path,
        flush=True
    )

    result = subprocess.run(
        [
            "exactdoc",
            input_path,
            "-o",
            output_path,
        ],
        check=True,
        timeout=300,
        capture_output=True,
        text=True,
    )

    if result.stdout:
        print(
            "ExactDoc output:",
            result.stdout,
            flush=True
        )

    if result.stderr:
        print(
            "ExactDoc warnings:",
            result.stderr,
            flush=True
        )

    return result


# ============================================================
# PDF / FORM DETECTION
# ============================================================

def pdf_widget_count(
    input_path
):

    pdf = None

    try:

        pdf = fitz.open(
            input_path
        )

        count = 0

        for page in pdf:

            widgets = page.widgets()

            if widgets:

                count += len(
                    list(widgets)
                )

        print(
            "Interactive widget count:",
            count,
            flush=True
        )

        return count

    except Exception as error:

        print(
            "Form detection error:",
            repr(error),
            flush=True
        )

        return 0

    finally:

        if pdf is not None:
            pdf.close()


def is_la350_form(
    input_path
):

    pdf = None

    try:

        pdf = fitz.open(
            input_path
        )

        if pdf.page_count != 1:
            return False

        text = (
            pdf[0].get_text("text")
            or ""
        )

        normalized = re.sub(
            r"\s+",
            " ",
            text
        ).lower()

        return (
            "la-350" in normalized
            and
            "notice of available language"
            in normalized
            and
            "service provider"
            in normalized
        )

    except Exception as error:

        print(
            "LA-350 detection error:",
            repr(error),
            flush=True
        )

        return False

    finally:

        if pdf is not None:
            pdf.close()


# ============================================================
# WORD XML HELPERS
# ============================================================

def set_cell_border(
    cell,
    top=None,
    bottom=None,
    left=None,
    right=None
):

    tc = cell._tc

    tcPr = tc.get_or_add_tcPr()

    tcBorders = (
        tcPr.first_child_found_in(
            "w:tcBorders"
        )
    )

    if tcBorders is None:

        tcBorders = OxmlElement(
            "w:tcBorders"
        )

        tcPr.append(
            tcBorders
        )

    for edge_name, edge in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():

        if edge is None:
            continue

        tag = (
            "w:"
            + edge_name
        )

        element = tcBorders.find(
            qn(tag)
        )

        if element is None:

            element = OxmlElement(
                tag
            )

            tcBorders.append(
                element
            )

        element.set(
            qn("w:val"),
            edge.get(
                "val",
                "single"
            )
        )

        element.set(
            qn("w:sz"),
            str(
                edge.get(
                    "sz",
                    6
                )
            )
        )

        element.set(
            qn("w:color"),
            edge.get(
                "color",
                "000000"
            )
        )


def set_cell_margins(
    cell,
    top=0,
    start=20,
    bottom=0,
    end=20
):

    tcPr = (
        cell._tc.get_or_add_tcPr()
    )

    tcMar = (
        tcPr.first_child_found_in(
            "w:tcMar"
        )
    )

    if tcMar is None:

        tcMar = OxmlElement(
            "w:tcMar"
        )

        tcPr.append(
            tcMar
        )

    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():

        node = tcMar.find(
            qn(
                "w:"
                + name
            )
        )

        if node is None:

            node = OxmlElement(
                "w:"
                + name
            )

            tcMar.append(
                node
            )

        node.set(
            qn("w:w"),
            str(value)
        )

        node.set(
            qn("w:type"),
            "dxa"
        )


def set_table_fixed_layout(
    table
):

    tblPr = table._tbl.tblPr

    tblLayout = (
        tblPr.first_child_found_in(
            "w:tblLayout"
        )
    )

    if tblLayout is None:

        tblLayout = OxmlElement(
            "w:tblLayout"
        )

        tblPr.append(
            tblLayout
        )

    tblLayout.set(
        qn("w:type"),
        "fixed"
    )


def set_cell_width(
    cell,
    width_inches
):

    cell.width = Inches(
        width_inches
    )

    tcPr = (
        cell._tc.get_or_add_tcPr()
    )

    tcW = (
        tcPr.first_child_found_in(
            "w:tcW"
        )
    )

    if tcW is None:

        tcW = OxmlElement(
            "w:tcW"
        )

        tcPr.append(
            tcW
        )

    tcW.set(
        qn("w:w"),
        str(
            int(
                width_inches
                * 1440
            )
        )
    )

    tcW.set(
        qn("w:type"),
        "dxa"
    )


def set_table_column_widths(
    table,
    widths
):

    grid = table._tbl.tblGrid

    grid_cols = grid.findall(
        qn("w:gridCol")
    )

    for index, width in enumerate(
        widths
    ):

        if index < len(
            grid_cols
        ):

            grid_cols[index].set(
                qn("w:w"),
                str(
                    int(
                        width
                        * 1440
                    )
                )
            )

    for row in table.rows:

        for index, width in enumerate(
            widths
        ):

            if index < len(
                row.cells
            ):

                set_cell_width(
                    row.cells[index],
                    width
                )


def set_row_height(
    row,
    points,
    exact=False
):

    trPr = (
        row._tr.get_or_add_trPr()
    )

    trHeight = OxmlElement(
        "w:trHeight"
    )

    trHeight.set(
        qn("w:val"),
        str(
            int(
                points
                * 20
            )
        )
    )

    trHeight.set(
        qn("w:hRule"),
        (
            "exact"
            if exact
            else "atLeast"
        )
    )

    trPr.append(
        trHeight
    )


def keep_table_row_together(
    row
):

    trPr = (
        row._tr.get_or_add_trPr()
    )

    if trPr.find(
        qn("w:cantSplit")
    ) is None:

        trPr.append(
            OxmlElement(
                "w:cantSplit"
            )
        )


def set_cell_shading(
    cell,
    fill
):

    tcPr = (
        cell._tc.get_or_add_tcPr()
    )

    shd = tcPr.find(
        qn("w:shd")
    )

    if shd is None:

        shd = OxmlElement(
            "w:shd"
        )

        tcPr.append(
            shd
        )

    shd.set(
        qn("w:fill"),
        fill
    )


def set_table_borders(
    table,
    size=5
):

    border = {
        "val": "single",
        "sz": size,
        "color": "000000"
    }

    for row in table.rows:

        for cell in row.cells:

            set_cell_border(
                cell,
                top=border,
                bottom=border,
                left=border,
                right=border
            )


def remove_cell_borders(
    cell
):

    none = {
        "val": "nil",
        "sz": 0,
        "color": "FFFFFF"
    }

    set_cell_border(
        cell,
        top=none,
        bottom=none,
        left=none,
        right=none
    )


def set_run_font(
    run,
    size=8,
    bold=False,
    italic=False,
    name="Arial"
):

    run.font.name = name

    run.font.size = Pt(
        size
    )

    run.bold = bold

    run.italic = italic

    rPr = (
        run._element.get_or_add_rPr()
    )

    rFonts = rPr.rFonts

    if rFonts is None:

        rFonts = OxmlElement(
            "w:rFonts"
        )

        rPr.append(
            rFonts
        )

    rFonts.set(
        qn("w:ascii"),
        name
    )

    rFonts.set(
        qn("w:hAnsi"),
        name
    )


def add_run(
    paragraph,
    text,
    size=8,
    bold=False,
    italic=False,
    color=None
):

    run = paragraph.add_run(
        text
    )

    set_run_font(
        run,
        size=size,
        bold=bold,
        italic=italic
    )

    if color:

        run.font.color.rgb = (
            RGBColor(
                *color
            )
        )

    return run


def clear_cell(
    cell
):

    cell.text = ""

    p = cell.paragraphs[0]

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    p.paragraph_format.line_spacing = 1.0

    return p


def set_normal_document_defaults(
    document
):

    normal = document.styles[
        "Normal"
    ]

    normal.font.name = "Arial"

    normal.font.size = Pt(
        7.3
    )

    normal.paragraph_format.space_before = Pt(
        0
    )

    normal.paragraph_format.space_after = Pt(
        0
    )

    normal.paragraph_format.line_spacing = 1.0


# ============================================================
# LA-350 FIELD EXTRACTION
# ============================================================

def widget_short_name(
    full_name
):

    if not full_name:
        return ""

    last = full_name.split(
        "."
    )[-1]

    return re.sub(
        r"\[\d+\]$",
        "",
        last
    )


def checkbox_is_checked(
    value
):

    if value is None:
        return False

    return str(
        value
    ).strip().lower() not in {
        "",
        "off",
        "0",
        "false",
        "none",
        "null",
        "no",
    }


def extract_la350_values(
    input_path
):

    values = {
        "date": "",
        "court": "",
        "typed_name": "",
        "provider": "",
        "address": "",
        "telephone": "",
        "contact_name": "",
        "email": "",
        "web": "",
        "calendar_year": "",
        "service_specify": "",
        "language_specify": "",
        "assistance_specify": "",
        "service_area": "",
        "narrative": False,
        "services": [False] * 10,
        "languages": [False] * 12,
        "assistance": [False] * 5,
    }

    pdf = fitz.open(
        input_path
    )

    try:

        page = pdf[0]

        widgets = list(
            page.widgets()
            or []
        )

        text_map = {
            "DateTimeField1":
                "date",
            "CourtInfo_ft":
                "court",
            "YourName_ft":
                "typed_name",
            "TextField3":
                "provider",
            "TextField4":
                "address",
            "TextField5":
                "telephone",
            "TextField7":
                "contact_name",
            "TextField8":
                "email",
            "TextField20":
                "web",
            "TextField21":
                "service_specify",
            "TextField22":
                "language_specify",
            "TextField23":
                "assistance_specify",
            "TextField24":
                "service_area",
            "TextField25":
                "calendar_year",
        }

        for widget in widgets:

            full_name = (
                widget.field_name
                or ""
            )

            short = widget_short_name(
                full_name
            )

            value = (
                widget.field_value
                or ""
            )

            if short in text_map:

                values[
                    text_map[short]
                ] = str(
                    value
                ).strip()

                continue

            if short == "CheckBox29":

                values[
                    "narrative"
                ] = checkbox_is_checked(
                    value
                )

                continue

            if (
                "Table1"
                in full_name
                and
                short.startswith(
                    "CheckBox"
                )
            ):

                match = re.search(
                    r"CheckBox(\d+)",
                    short
                )

                if match:

                    number = int(
                        match.group(1)
                    )

                    mapping = {
                        1: 0,
                        2: 1,
                        3: 2,
                        4: 3,
                        5: 4,
                        6: 5,
                        7: 6,
                        8: 7,
                        9: 8,
                        28: 9,
                    }

                    index = mapping.get(
                        number
                    )

                    if index is not None:

                        values[
                            "services"
                        ][
                            index
                        ] = checkbox_is_checked(
                            value
                        )

                continue

            if (
                "Table2"
                in full_name
                and
                short.startswith(
                    "CheckBox"
                )
            ):

                match = re.search(
                    r"CheckBox(\d+)",
                    short
                )

                if match:

                    number = int(
                        match.group(1)
                    )

                    if 11 <= number <= 22:

                        values[
                            "languages"
                        ][
                            number - 11
                        ] = checkbox_is_checked(
                            value
                        )

                continue

            if (
                "Table3"
                in full_name
                and
                short.startswith(
                    "CheckBox"
                )
            ):

                match = re.search(
                    r"CheckBox(\d+)",
                    short
                )

                if match:

                    number = int(
                        match.group(1)
                    )

                    if 23 <= number <= 27:

                        values[
                            "assistance"
                        ][
                            number - 23
                        ] = checkbox_is_checked(
                            value
                        )

                continue

        return values

    finally:

        pdf.close()


# ============================================================
# LA-350 FORM HELPERS
# ============================================================

def checkbox_symbol(
    checked
):

    return (
        "☒"
        if checked
        else "☐"
    )


def add_form_line(
    paragraph,
    label,
    value="",
    width_chars=30,
    size=7.0
):

    add_run(
        paragraph,
        label,
        size=size
    )

    add_run(
        paragraph,
        " "
        + (
            value
            if value
            else "_" * width_chars
        ),
        size=size
    )


def add_checkbox_line(
    cell,
    label,
    checked=False,
    size=6.9
):

    p = clear_cell(
        cell
    )

    add_run(
        p,
        checkbox_symbol(
            checked
        )
        + " ",
        size=size
    )

    add_run(
        p,
        label,
        size=size
    )

    return p


# ============================================================
# LA-350 HEADER
# ============================================================

def add_la350_header(
    document
):

    table = document.add_table(
        rows=1,
        cols=3
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = False

    set_table_fixed_layout(
        table
    )

    # Slightly narrower LA-350 black block.
    # Slightly wider center title.
    widths = [
        2.30,
        2.80,
        2.40
    ]

    set_table_column_widths(
        table,
        widths
    )

    for index, width in enumerate(
        widths
    ):

        set_cell_width(
            table.cell(
                0,
                index
            ),
            width
        )

        set_cell_margins(
            table.cell(
                0,
                index
            ),
            top=10,
            bottom=8,
            start=28,
            end=28
        )

    # BLACK LA-350 BOX

    left = table.cell(
        0,
        0
    )

    set_cell_shading(
        left,
        "000000"
    )

    remove_cell_borders(
        left
    )

    p = clear_cell(
        left
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    add_run(
        p,
        "LA-350",
        size=14.0,
        bold=True,
        color=(
            255,
            255,
            255
        )
    )

    # TITLE

    middle = table.cell(
        0,
        1
    )

    remove_cell_borders(
        middle
    )

    p = clear_cell(
        middle
    )

    add_run(
        p,
        "Notice of Available Language\n"
        "Assistance—Service Provider",
        size=11.8,
        bold=True
    )

    # CLERK BOX

    right = table.cell(
        0,
        2
    )

    p = clear_cell(
        right
    )

    add_run(
        p,
        "Clerk stamps date here when form is received.",
        size=5.7,
        italic=True
    )

    border = {
        "val": "single",
        "sz": 6,
        "color": "000000"
    }

    set_cell_border(
        right,
        top=border,
        bottom=border,
        left=border,
        right=border
    )

    # Slightly shorter than test (14)
    set_row_height(
        table.rows[0],
        48,
        exact=True
    )

    keep_table_row_together(
        table.rows[0]
    )

    return table


# ============================================================
# INTRO / COURT INFO
# ============================================================

def add_la350_intro(
    document,
    values
):

    table = document.add_table(
        rows=2,
        cols=2
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = False

    set_table_fixed_layout(
        table
    )

    widths = [
        4.90,
        2.60
    ]

    set_table_column_widths(
        table,
        widths
    )

    for row_index in range(
        2
    ):

        for col_index in range(
            2
        ):

            set_cell_width(
                table.cell(
                    row_index,
                    col_index
                ),
                widths[
                    col_index
                ]
            )

            set_cell_margins(
                table.cell(
                    row_index,
                    col_index
                ),
                top=2,
                bottom=1,
                start=18,
                end=24
            )

            remove_cell_borders(
                table.cell(
                    row_index,
                    col_index
                )
            )

    # USE THIS FORM TO

    left = table.cell(
        0,
        0
    )

    p = clear_cell(
        left
    )

    add_run(
        p,
        "Use this form to:",
        size=7.0,
        bold=True
    )

    p = left.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    p.paragraph_format.line_spacing = 0.90

    add_run(
        p,
        "• Tell the court that you are a service provider, "
        "program, or professional offering language assistance "
        "with services that may be ordered by a court; and",
        size=6.7
    )

    p = left.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    p.paragraph_format.line_spacing = 0.90

    add_run(
        p,
        "• Provide information about the services you provide, "
        "the languages and types of language assistance available, "
        "and your service area.",
        size=6.7
    )

    # MERGED COURT BOX

    right = table.cell(
        0,
        1
    ).merge(
        table.cell(
            1,
            1
        )
    )

    p = clear_cell(
        right
    )

    add_run(
        p,
        "Fill in court name and address:",
        size=5.7,
        italic=True
    )

    p = right.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    add_run(
        p,
        "Superior Court of California, County of",
        size=6.6,
        bold=True
    )

    p = right.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    add_run(
        p,
        (
            values[
                "court"
            ]
            if values[
                "court"
            ]
            else " "
        ),
        size=6.7
    )

    border = {
        "val": "single",
        "sz": 6,
        "color": "000000"
    }

    set_cell_border(
        right,
        top=border,
        bottom=border,
        left=border,
        right=border
    )

    # SECTION 1

    left2 = table.cell(
        1,
        0
    )

    p = clear_cell(
        left2
    )

    add_run(
        p,
        "1  ",
        size=7.3,
        bold=True
    )

    add_run(
        p,
        "This form should be filed with the court by January 31 "
        "of each year to indicate services that will be provided "
        "during the calendar year. You may also submit this form "
        "to let the court know your services have changed.",
        size=6.6
    )

    p = left2.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    p.paragraph_format.line_spacing = 0.92

    add_run(
        p,
        "    The information in this form describes services "
        "available during calendar year: ",
        size=6.6
    )

    add_run(
        p,
        (
            values[
                "calendar_year"
            ]
            if values[
                "calendar_year"
            ]
            else "____________"
        ),
        size=6.6
    )

    # Reduced from 86 / 58
    set_row_height(
        table.rows[0],
        78,
        exact=True
    )

    set_row_height(
        table.rows[1],
        50,
        exact=True
    )

    keep_table_row_together(
        table.rows[0]
    )

    keep_table_row_together(
        table.rows[1]
    )

    return table


# ============================================================
# PROVIDER SECTION
# ============================================================

def add_la350_provider_section(
    document,
    values
):

    table = document.add_table(
        rows=4,
        cols=2
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = False

    set_table_fixed_layout(
        table
    )

    widths = [
        0.33,
        7.17
    ]

    set_table_column_widths(
        table,
        widths
    )

    for row_index in range(
        4
    ):

        set_cell_width(
            table.cell(
                row_index,
                0
            ),
            widths[0]
        )

        set_cell_width(
            table.cell(
                row_index,
                1
            ),
            widths[1]
        )

        remove_cell_borders(
            table.cell(
                row_index,
                0
            )
        )

        remove_cell_borders(
            table.cell(
                row_index,
                1
            )
        )

        set_cell_margins(
            table.cell(
                row_index,
                0
            ),
            top=0,
            bottom=0,
            start=0,
            end=0
        )

        set_cell_margins(
            table.cell(
                row_index,
                1
            ),
            top=0,
            bottom=0,
            start=8,
            end=8
        )

    p = clear_cell(
        table.cell(
            0,
            0
        )
    )

    add_run(
        p,
        "2",
        size=7.5,
        bold=True
    )

    p = clear_cell(
        table.cell(
            0,
            1
        )
    )

    add_form_line(
        p,
        "Name of service provider:",
        values["provider"],
        width_chars=46,
        size=6.8
    )

    p = clear_cell(
        table.cell(
            1,
            1
        )
    )

    add_form_line(
        p,
        "Address:",
        values["address"],
        width_chars=78,
        size=6.8
    )

    p = clear_cell(
        table.cell(
            2,
            1
        )
    )

    add_form_line(
        p,
        "Telephone:",
        values["telephone"],
        width_chars=18,
        size=6.8
    )

    add_run(
        p,
        "      ",
        size=6.8
    )

    add_form_line(
        p,
        "Web address:",
        values["web"],
        width_chars=28,
        size=6.8
    )

    p = clear_cell(
        table.cell(
            3,
            1
        )
    )

    add_form_line(
        p,
        "Contact name:",
        values["contact_name"],
        width_chars=25,
        size=6.8
    )

    add_run(
        p,
        "      ",
        size=6.8
    )

    add_form_line(
        p,
        "E-mail:",
        values["email"],
        width_chars=30,
        size=6.8
    )

    # Slightly tighter than test (14)
    for row_index, height in enumerate(
        [
            16,
            16,
            16,
            16
        ]
    ):

        set_row_height(
            table.rows[
                row_index
            ],
            height,
            exact=True
        )

        keep_table_row_together(
            table.rows[
                row_index
            ]
        )

    return table


# ============================================================
# OPTION TABLE
# ============================================================

def add_option_table(
    parent_cell,
    title,
    labels,
    checked_values,
    specify_value="",
    include_service_area=False,
    service_area_value=""
):

    rows_needed = (
        1
        + len(labels)
        + 1
        + (
            1
            if include_service_area
            else 0
        )
    )

    table = parent_cell.add_table(
        rows=rows_needed,
        cols=1
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = False

    set_table_fixed_layout(
        table
    )

    set_table_borders(
        table,
        size=4
    )

    # HEADER

    header = table.cell(
        0,
        0
    )

    p = clear_cell(
        header
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    add_run(
        p,
        title,
        size=7.6,
        bold=True
    )

    p2 = header.add_paragraph()

    p2.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p2.paragraph_format.space_before = Pt(
        0
    )

    p2.paragraph_format.space_after = Pt(
        0
    )

    add_run(
        p2,
        "(select all that apply)",
        size=5.9,
        italic=True
    )

    set_cell_margins(
        header,
        top=7,
        bottom=7,
        start=18,
        end=18
    )

    set_row_height(
        table.rows[0],
        31,
        exact=True
    )

    # OPTION ROWS

    for index, label in enumerate(
        labels
    ):

        cell = table.cell(
            index + 1,
            0
        )

        add_checkbox_line(
            cell,
            label,
            checked_values[
                index
            ],
            size=6.4
        )

        set_cell_margins(
            cell,
            top=2,
            bottom=2,
            start=18,
            end=18
        )

        if len(label) > 34:

            height = 33

        elif len(label) > 27:

            height = 28

        else:

            height = 23

        set_row_height(
            table.rows[
                index + 1
            ],
            height,
            exact=True
        )

        keep_table_row_together(
            table.rows[
                index + 1
            ]
        )

    # SPECIFY

    specify_row = (
        1
        + len(labels)
    )

    cell = table.cell(
        specify_row,
        0
    )

    p = clear_cell(
        cell
    )

    add_form_line(
        p,
        "Specify:",
        specify_value,
        width_chars=17,
        size=6.4
    )

    set_cell_margins(
        cell,
        top=2,
        bottom=2,
        start=18,
        end=18
    )

    set_row_height(
        table.rows[
            specify_row
        ],
        23,
        exact=True
    )

    # SERVICE AREA

    if include_service_area:

        index = (
            specify_row
            + 1
        )

        cell = table.cell(
            index,
            0
        )

        p = clear_cell(
            cell
        )

        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        add_run(
            p,
            "Service Area",
            size=7.5,
            bold=True
        )

        p2 = cell.add_paragraph()

        p2.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        p2.paragraph_format.space_before = Pt(
            0
        )

        p2.paragraph_format.space_after = Pt(
            0
        )

        add_run(
            p2,
            "(county or region)",
            size=5.9
        )

        p3 = cell.add_paragraph()

        p3.paragraph_format.space_before = Pt(
            2
        )

        p3.paragraph_format.space_after = Pt(
            0
        )

        add_run(
            p3,
            (
                service_area_value
                if service_area_value
                else "\n________________________"
            ),
            size=6.4
        )

        set_cell_margins(
            cell,
            top=6,
            bottom=6,
            start=20,
            end=20
        )

        set_row_height(
            table.rows[
                index
            ],
            96,
            exact=True
        )

    return table


# ============================================================
# SECTION 3 / THREE LOWER TABLES
# ============================================================

def add_la350_services_section(
    document,
    values
):

    heading = document.add_table(
        rows=1,
        cols=2
    )

    heading.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    heading.autofit = False

    set_table_fixed_layout(
        heading
    )

    heading_widths = [
        3.15,
        4.35
    ]

    set_table_column_widths(
        heading,
        heading_widths
    )

    for index, width in enumerate(
        heading_widths
    ):

        set_cell_width(
            heading.cell(
                0,
                index
            ),
            width
        )

    for cell in heading.rows[
        0
    ].cells:

        remove_cell_borders(
            cell
        )

        set_cell_margins(
            cell,
            top=0,
            bottom=0,
            start=4,
            end=4
        )

    p = clear_cell(
        heading.cell(
            0,
            0
        )
    )

    add_run(
        p,
        "3  Information about the services provided:",
        size=6.8
    )

    p = clear_cell(
        heading.cell(
            0,
            1
        )
    )

    add_run(
        p,
        checkbox_symbol(
            values[
                "narrative"
            ]
        )
        + " ",
        size=6.7
    )

    add_run(
        p,
        "Check here to attach a narrative description "
        "of the services offered.",
        size=6.4
    )

    set_row_height(
        heading.rows[0],
        20,
        exact=True
    )

    keep_table_row_together(
        heading.rows[0]
    )

    # --------------------------------------------------------
    # FIVE-COLUMN OUTER TABLE
    #
    # 0 = Services
    # 1 = spacer
    # 2 = Languages
    # 3 = spacer
    # 4 = Assistance
    #
    # This creates visible gaps between the three boxes.
    # --------------------------------------------------------

    outer = document.add_table(
        rows=1,
        cols=5
    )

    outer.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    outer.autofit = False

    set_table_fixed_layout(
        outer
    )

    widths = [
        2.55,
        0.14,
        1.92,
        0.14,
        2.75
    ]

    set_table_column_widths(
        outer,
        widths
    )

    for index, width in enumerate(
        widths
    ):

        cell = outer.cell(
            0,
            index
        )

        set_cell_width(
            cell,
            width
        )

        remove_cell_borders(
            cell
        )

        set_cell_margins(
            cell,
            top=0,
            bottom=0,
            start=0,
            end=0
        )

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.TOP
        )

        clear_cell(
            cell
        )

    # Spacer cells stay empty.
    outer.cell(
        0,
        1
    ).text = ""

    outer.cell(
        0,
        3
    ).text = ""

    service_labels = [
        "Mediation",
        "Child custody recommending counseling",
        "Professional supervised child visitation",
        "Parenting education classes",
        "Anger management classes",
        "Mental health counseling",
        "Batterer intervention–MEN",
        "Batterer intervention–WOMEN",
        "Alcohol/substance abuse treatment",
        "Other",
    ]

    language_labels = [
        "Any language",
        "American Sign Language",
        "Spanish",
        "Mandarin",
        "Cantonese",
        "Farsi",
        "Korean",
        "Punjabi",
        "Russian",
        "Tagalog",
        "Vietnamese",
        "Other",
    ]

    assistance_labels = [
        "Program offered directly in language",
        "In-person interpreter",
        "Telephone interpreter",
        "Translated materials",
        "Other",
    ]

    add_option_table(
        outer.cell(
            0,
            0
        ),
        "Services",
        service_labels,
        values[
            "services"
        ],
        specify_value=values[
            "service_specify"
        ]
    )

    add_option_table(
        outer.cell(
            0,
            2
        ),
        "Languages Available",
        language_labels,
        values[
            "languages"
        ],
        specify_value=values[
            "language_specify"
        ]
    )

    add_option_table(
        outer.cell(
            0,
            4
        ),
        "Types of Language\nAssistance",
        assistance_labels,
        values[
            "assistance"
        ],
        specify_value=values[
            "assistance_specify"
        ],
        include_service_area=True,
        service_area_value=values[
            "service_area"
        ]
    )

    keep_table_row_together(
        outer.rows[0]
    )

    return outer


# ============================================================
# SIGNATURE AND FOOTER
# ============================================================

def add_la350_signature_and_footer(
    document,
    values
):

    p = document.add_paragraph()

    p.paragraph_format.space_before = Pt(
        0
    )

    p.paragraph_format.space_after = Pt(
        0
    )

    p.paragraph_format.line_spacing = 0.9

    add_form_line(
        p,
        "Date:",
        values[
            "date"
        ],
        width_chars=20,
        size=6.6
    )

    table = document.add_table(
        rows=2,
        cols=2
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.autofit = False

    set_table_fixed_layout(
        table
    )

    signature_widths = [
        3.75,
        3.75
    ]

    set_table_column_widths(
        table,
        signature_widths
    )

    for row_index in range(
        2
    ):

        for col_index in range(
            2
        ):

            set_cell_width(
                table.cell(
                    row_index,
                    col_index
                ),
                3.75
            )

            remove_cell_borders(
                table.cell(
                    row_index,
                    col_index
                )
            )

            set_cell_margins(
                table.cell(
                    row_index,
                    col_index
                ),
                top=0,
                bottom=0,
                start=0,
                end=0
            )

    p = clear_cell(
        table.cell(
            0,
            0
        )
    )

    add_run(
        p,
        (
            values[
                "typed_name"
            ]
            if values[
                "typed_name"
            ]
            else
            "________________________________________"
        ),
        size=6.5
    )

    p = clear_cell(
        table.cell(
            0,
            1
        )
    )

    add_run(
        p,
        "________________________________________",
        size=6.5
    )

    p = clear_cell(
        table.cell(
            1,
            0
        )
    )

    add_run(
        p,
        "Type or print your name",
        size=5.9,
        italic=True
    )

    p = clear_cell(
        table.cell(
            1,
            1
        )
    )

    add_run(
        p,
        "Sign your name",
        size=5.9,
        italic=True
    )

    set_row_height(
        table.rows[0],
        9,
        exact=True
    )

    set_row_height(
        table.rows[1],
        8,
        exact=True
    )

    keep_table_row_together(
        table.rows[0]
    )

    keep_table_row_together(
        table.rows[1]
    )

    # FOOTER

    footer = document.add_table(
        rows=1,
        cols=3
    )

    footer.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    footer.autofit = False

    set_table_fixed_layout(
        footer
    )

    widths = [
        2.35,
        3.50,
        1.65
    ]

    set_table_column_widths(
        footer,
        widths
    )

    for index, width in enumerate(
        widths
    ):

        set_cell_width(
            footer.cell(
                0,
                index
            ),
            width
        )

        remove_cell_borders(
            footer.cell(
                0,
                index
            )
        )

        set_cell_margins(
            footer.cell(
                0,
                index
            ),
            top=0,
            bottom=0,
            start=0,
            end=0
        )

    p = clear_cell(
        footer.cell(
            0,
            0
        )
    )

    add_run(
        p,
        "Judicial Council of California, www.courts.ca.gov\n"
        "New September 1, 2019, Optional Form\n"
        "Cal. Rules of Court, rule 1.300",
        size=4.6
    )

    p = clear_cell(
        footer.cell(
            0,
            1
        )
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    add_run(
        p,
        "Notice of Available Language\n"
        "Assistance—Service Provider",
        size=8.5,
        bold=True
    )

    p = clear_cell(
        footer.cell(
            0,
            2
        )
    )

    p.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    add_run(
        p,
        "LA-350, Page 1 of 1",
        size=5.4,
        bold=True
    )

    # Slightly shorter footer pulls content upward.
    set_row_height(
        footer.rows[0],
        30,
        exact=True
    )

    keep_table_row_together(
        footer.rows[0]
    )


# ============================================================
# LA-350 -> DOCX
# ============================================================

def convert_la350_to_docx(
    input_path,
    output_path
):

    print(
        "LA-350 detected.",
        flush=True
    )

    print(
        "Using dedicated native Word LA-350 reconstruction.",
        flush=True
    )

    values = extract_la350_values(
        input_path
    )

    document = Document()

    set_normal_document_defaults(
        document
    )

    section = document.sections[
        0
    ]

    section.page_width = Inches(
        8.5
    )

    section.page_height = Inches(
        11
    )

    section.top_margin = Inches(
        0.08
    )

    section.bottom_margin = Inches(
        0.08
    )

    section.left_margin = Inches(
        0.15
    )

    section.right_margin = Inches(
        0.15
    )

    section.header_distance = Inches(
        0
    )

    section.footer_distance = Inches(
        0
    )

    if document.paragraphs:

        first = document.paragraphs[
            0
        ]

        first.paragraph_format.space_before = Pt(
            0
        )

        first.paragraph_format.space_after = Pt(
            0
        )

        first.paragraph_format.line_spacing = 0.1

        add_run(
            first,
            "",
            size=1
        )

    add_la350_header(
        document
    )

    add_la350_intro(
        document,
        values
    )

    add_la350_provider_section(
        document,
        values
    )

    add_la350_services_section(
        document,
        values
    )

    add_la350_signature_and_footer(
        document,
        values
    )

    document.save(
        output_path
    )

    print(
        "LA-350 native DOCX created:",
        output_path,
        flush=True
    )


# ============================================================
# UNKNOWN INTERACTIVE FORM FALLBACK
# ============================================================

def convert_interactive_form_image_fallback(
    input_path,
    output_path
):

    print(
        "Unknown interactive form.",
        flush=True
    )

    print(
        "Using safe image-preserving Word fallback.",
        flush=True
    )

    pdf = fitz.open(
        input_path
    )

    document = Document()

    set_normal_document_defaults(
        document
    )

    try:

        for page_index in range(
            pdf.page_count
        ):

            page = pdf[
                page_index
            ]

            if page_index > 0:

                document.add_page_break()

            section = document.sections[
                -1
            ]

            section.page_width = Pt(
                page.rect.width
            )

            section.page_height = Pt(
                page.rect.height
            )

            section.top_margin = Pt(
                6
            )

            section.bottom_margin = Pt(
                6
            )

            section.left_margin = Pt(
                6
            )

            section.right_margin = Pt(
                6
            )

            section.header_distance = Pt(
                0
            )

            section.footer_distance = Pt(
                0
            )

            pix = page.get_pixmap(
                matrix=fitz.Matrix(
                    2,
                    2
                ),
                alpha=False,
                annots=True
            )

            image_path = os.path.join(
                os.path.dirname(
                    output_path
                ),
                (
                    f"interactive_page_"
                    f"{page_index + 1}.png"
                )
            )

            pix.save(
                image_path
            )

            paragraph = (
                document.add_paragraph()
            )

            paragraph.paragraph_format.space_before = Pt(
                0
            )

            paragraph.paragraph_format.space_after = Pt(
                0
            )

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            run = paragraph.add_run()

            run.add_picture(
                image_path,
                width=Pt(
                    max(
                        page.rect.width
                        - 12,
                        72
                    )
                )
            )

        document.save(
            output_path
        )

    finally:

        pdf.close()

    print(
        "Interactive-form fallback DOCX created:",
        output_path,
        flush=True
    )


# ============================================================
# PDF -> WORD
# ============================================================

@app.route(
    "/convert/pdf-to-word",
    methods=["POST"]
)
def pdf_to_word():

    if "file" not in request.files:

        return jsonify({
            "error":
            "No file uploaded."
        }), 400

    uploaded_file = (
        request.files["file"]
    )

    if uploaded_file.filename == "":

        return jsonify({
            "error":
            "No file selected."
        }), 400

    filename = secure_filename(
        uploaded_file.filename
    )

    if not filename.lower().endswith(
        ".pdf"
    ):

        return jsonify({
            "error":
            "Please upload a PDF file."
        }), 400

    with tempfile.TemporaryDirectory() as temp_dir:

        input_path = os.path.join(
            temp_dir,
            filename
        )

        uploaded_file.save(
            input_path
        )

        output_name = (
            os.path.splitext(
                filename
            )[0]
            + ".docx"
        )

        output_path = os.path.join(
            temp_dir,
            output_name
        )

        try:

            print(
                "=================================",
                flush=True
            )

            print(
                "PDF TO WORD START:",
                filename,
                flush=True
            )

            print(
                "=================================",
                flush=True
            )

            widget_count = pdf_widget_count(
                input_path
            )

            if widget_count > 0:

                if is_la350_form(
                    input_path
                ):

                    convert_la350_to_docx(
                        input_path,
                        output_path
                    )

                else:

                    convert_interactive_form_image_fallback(
                        input_path,
                        output_path
                    )

            else:

                print(
                    "Normal PDF detected.",
                    flush=True
                )

                print(
                    "Using ExactDoc.",
                    flush=True
                )

                run_exactdoc(
                    input_path,
                    output_path
                )

        except subprocess.TimeoutExpired:

            print(
                "PDF conversion timed out.",
                flush=True
            )

            return jsonify({
                "error":
                "The PDF conversion took too long."
            }), 504

        except subprocess.CalledProcessError as error:

            print(
                "PDF converter exit code:",
                error.returncode,
                flush=True
            )

            print(
                "PDF converter stdout:",
                error.stdout,
                flush=True
            )

            print(
                "PDF converter stderr:",
                error.stderr,
                flush=True
            )

            return jsonify({
                "error":
                "The PDF could not be converted to Word."
            }), 500

        except Exception as error:

            print(
                "PDF to Word error:",
                repr(error),
                flush=True
            )

            return jsonify({
                "error":
                "The PDF could not be converted to Word."
            }), 500

        if not os.path.exists(
            output_path
        ):

            return jsonify({
                "error":
                "The Word document could not be created."
            }), 500

        if os.path.getsize(
            output_path
        ) == 0:

            return jsonify({
                "error":
                "The Word document was created but was empty."
            }), 500

        print(
            "=================================",
            flush=True
        )

        print(
            "PDF TO WORD SUCCESS:",
            output_name,
            flush=True
        )

        print(
            "=================================",
            flush=True
        )

        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_name,
            mimetype=(
                "application/vnd."
                "openxmlformats-officedocument."
                "wordprocessingml.document"
            )
        )


# ============================================================
# START SERVER
# ============================================================

if __name__ == "__main__":

    port = int(
        os.environ.get(
            "PORT",
            10000
        )
    )

    app.run(
        host="0.0.0.0",
        port=port
    )
